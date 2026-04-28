#!/usr/bin/env python
"""
Verify that the styled paid-media section renders into the DOCX when the
finance entry of structured_data contains a `paid_media` dict.

Asserts:
  1. _strip_inline_paid_media removes the `### Paid Media Plan` block
     between sibling H3 headings (without disturbing the rest).
  2. _strip_inline_paid_media stops at the next `## ` H2 boundary.
  3. _paid_media_digital_table builds the expected headers + a row per channel.
  4. _render_paid_media_section adds a Heading 3 'Paid Media Plan' to the doc.
  5. _render_paid_media_section adds at least one styled table (the digital
     channels rollup).
  6. _write_docx end-to-end emits a file that contains 'Paid Media Plan' as
     a heading and includes the digital channel labels.

Run from powerbuilder/ with: ./venv/bin/python scripts/_test_export_paid_media_docx.py
"""
import os
import sys
from pathlib import Path

import django

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
sys.path.insert(0, str(PROJECT_DIR))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "powerbuilder_app.settings")
django.setup()

from chat.agents import export as export_mod
from chat.agents.paid_media import estimate_paid_media


def main() -> None:
    # ---- 1. _strip_inline_paid_media: removes block between H3 siblings ----
    sample = (
        "## Budget Estimate\n"
        "Some budget narrative here.\n\n"
        "### Paid Media Plan\n"
        "Tier: **Tier 2 ($100K)**.\n\n"
        "**Digital channels**\n\n"
        "| Channel | Spend |\n"
        "|---|---|\n"
        "| Meta | $20,000 |\n\n"
        "### Win Number Calculation\n"
        "Win narrative.\n"
    )
    stripped = export_mod._strip_inline_paid_media(sample)
    assert "### Paid Media Plan" not in stripped, "H3 heading not stripped"
    assert "| Meta | $20,000 |" not in stripped, "table row not stripped"
    assert "### Win Number Calculation" in stripped, "next H3 should remain"
    assert "## Budget Estimate" in stripped, "parent H2 should remain"
    assert "Some budget narrative here." in stripped, "sibling text should remain"
    print("✓ 1. _strip_inline_paid_media removes block between H3 siblings")

    # ---- 2. _strip_inline_paid_media stops at next H2 ----
    sample_h2 = (
        "## Budget Estimate\n"
        "Budget prose.\n\n"
        "### Paid Media Plan\n"
        "Some content here.\n\n"
        "## Win Number Calculation\n"
        "Win prose.\n"
    )
    stripped_h2 = export_mod._strip_inline_paid_media(sample_h2)
    assert "### Paid Media Plan" not in stripped_h2
    assert "Some content here." not in stripped_h2
    assert "## Win Number Calculation" in stripped_h2
    assert "Win prose." in stripped_h2
    print("✓ 2. _strip_inline_paid_media stops at next H2 boundary")

    # ---- 3. _paid_media_digital_table: headers + rows ----
    estimate = estimate_paid_media(
        budget=100_000,
        query="paid media budget for the GOTV program",
        language_intent="en",
        district_label="HD-101 GA",
        target_universe=15_000,
        flight_weeks=6,
    )
    assert estimate is not None, "estimate_paid_media returned None"
    headers, rows = export_mod._paid_media_digital_table(estimate["channels"])
    assert headers[0] == "Channel"
    assert "Spend" in headers
    assert "Est. Lift (pp)" in headers
    assert len(rows) == len(estimate["channels"]), "row count must match channels"
    # Each row must have one cell per header.
    for row in rows:
        assert len(row) == len(headers), f"row width mismatch: {row}"
    print(f"✓ 3. _paid_media_digital_table built {len(rows)} rows / {len(headers)} cols")

    # ---- 4. _render_paid_media_section: adds 'Paid Media Plan' heading ----
    from docx import Document
    doc = Document()
    export_mod._render_paid_media_section(doc, estimate)
    headings = [
        p.text for p in doc.paragraphs
        if p.style and p.style.name and p.style.name.startswith("Heading")
    ]
    assert "Paid Media Plan" in headings, f"Paid Media Plan heading missing: {headings}"
    print("✓ 4. _render_paid_media_section added 'Paid Media Plan' heading")

    # ---- 5. _render_paid_media_section: at least one table ----
    assert len(doc.tables) >= 1, "expected at least one table in the section"
    digital_table = doc.tables[0]
    # First row is headers
    hdr_cells = [c.text for c in digital_table.rows[0].cells]
    assert "Channel" in hdr_cells
    assert "Spend" in hdr_cells
    print(f"✓ 5. _render_paid_media_section added {len(doc.tables)} table(s)")

    # ---- 6. End-to-end _write_docx with finance entry that has paid_media ----
    structured_data = [
        {
            "agent": "win_number",
            "win_number": 4_500,
            "projected_turnout": 22_000,
            "voter_universe_cvap": 38_000,
            "avg_turnout_pct": 0.58,
            "victory_margin": 0.04,
            "historical_context": "GA-07 2022 midterm baseline",
            "district_label": "HD-101 GA",
        },
        {
            "agent": "precincts",
            "precincts": [
                {"precinct_geoid": "13135-A1", "precinct_name": "P-A1",
                 "registered_voters": 2400, "lean_pct": 0.61},
                {"precinct_geoid": "13135-A2", "precinct_name": "P-A2",
                 "registered_voters": 2100, "lean_pct": 0.55},
            ],
        },
        {
            "agent": "finance",
            "budget": 100_000,
            "unit_costs": {
                "door_knock": 5.0, "phone_call": 1.5,
                "text_message": 0.30, "mail_piece": 0.85,
            },
            "budget_program": {
                "door_knock": {"unit_cost": 5.0, "contacts": 8000, "budget_allocated": 40_000},
                "phone_call": {"unit_cost": 1.5, "contacts": 10000, "budget_allocated": 15_000},
                "text_message": {"unit_cost": 0.30, "contacts": 50000, "budget_allocated": 15_000},
                "mail_piece": {"unit_cost": 0.85, "contacts": 35000, "budget_allocated": 30_000},
            },
            "paid_media": estimate,
        },
    ]
    synthesis = (
        "# HD-101 GA — Political Program Plan\n\n"
        "## Executive Summary\nWe will run a paid + field hybrid program.\n\n"
        "## District Background\nGwinnett-adjacent suburban district.\n\n"
        "## Target Universe and Demographics\nGrowing Latinx and AAPI share.\n\n"
        "## Geographic Targeting\nFocus on top precincts.\n\n"
        "## Messaging Strategy\nIn-language doors and SMS.\n\n"
        "## Budget Estimate\n"
        "We have a $100K total program budget.\n\n"
        "### Paid Media Plan\n"
        "This inline block should be stripped from the docx.\n\n"
        "| Channel | Spend |\n|---|---|\n| Meta | $20,000 |\n\n"
        "## Win Number Calculation\nThe win number is 4,500.\n\n"
        "## Program Recommendations\n- Lead with doors\n- Layer SMS\n"
    )
    state = {
        "query": "build a paid media plan for HD-101",
        "structured_data": structured_data,
        "active_agents": list(export_mod.PLAN_AGENTS),
        "research_results": [],
        "errors": [],
    }
    out = export_mod._write_docx(synthesis, state, "HD-101 GA")
    docx_path = out.get("generated_file_path")
    assert docx_path and os.path.exists(docx_path), f"docx not written: {docx_path}"

    # Re-open and confirm 'Paid Media Plan' heading is present and the inline
    # markdown table line did not leak through.
    written = Document(docx_path)
    written_headings = [
        p.text for p in written.paragraphs
        if p.style and p.style.name and p.style.name.startswith("Heading")
    ]
    written_text = "\n".join(p.text for p in written.paragraphs)
    assert "Paid Media Plan" in written_headings, \
        f"Paid Media Plan heading missing from docx: {written_headings}"
    assert "This inline block should be stripped from the docx." not in written_text, \
        "stripped markdown leaked into the docx"
    # And the digital channels rollup table should be present.
    found_channel_table = False
    for table in written.tables:
        hdrs = [c.text for c in table.rows[0].cells]
        if "Channel" in hdrs and "Est. Lift (pp)" in hdrs:
            found_channel_table = True
            break
    assert found_channel_table, "digital channel rollup table missing from docx"
    print(f"✓ 6. _write_docx end-to-end emitted styled paid-media section: {docx_path}")

    print("\nAll 6 assertions passed.")


if __name__ == "__main__":
    main()
