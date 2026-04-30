# powerbuilder/chat/agents/export.py
"""
Synthesizer and Reporter — the final node in the LangGraph graph.

Two jobs:

  1. SYNTHESIS: Reads all agent outputs from AgentState, deduplicates
     research_results by content hash, resolves contradictions by deferring
     to the most specific and most recently dated source (via the LLM system
     prompt), flags errors to the user, and calls GPT-4o to produce a clean
     non-repetitive professional program briefing in Markdown.

  2. FORMATTING: Writes the briefing to the requested output_format:
       - 'text' | 'markdown': writes synthesized Markdown directly to final_answer.
       - 'docx': builds a Word document via python-docx. Prose comes from the LLM
                 synthesis; data tables (win number, precincts, budget) are built
                 programmatically from structured_data for accuracy. Saves to
                 EXPORTS_DIR and writes path to generated_file_path.
       - 'xlsx': builds an Excel workbook via openpyxl with three sheets —
                 Precinct Targets, Win Number, Budget Estimate. Saves to EXPORTS_DIR.
       - 'csv':  flat CSV of target precincts (falls back to a structured_data
                 summary if no precinct data exists). Saves to EXPORTS_DIR.

System prompt (verbatim from spec):
  "You are a senior political strategist. You have received findings from specialist
  analysts. Synthesize them into a single clear non-repetitive professional program
  briefing. Do not invent information not present in the inputs."
"""

import csv as csv_module
import hashlib
import logging
import os
import re
from datetime import datetime
from typing import Optional

from dotenv import load_dotenv
load_dotenv()

from ..utils.llm_config import get_completion_client

from .state import AgentState

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

TEMPLATES_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "../../tool_templates")
)
# Word template used as base for docx output — inherits styles, margins, and footer.
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "political_plan_template.docx")

# Override with EXPORTS_DIR env var in production (e.g. Django MEDIA_ROOT subdir).
EXPORTS_DIR = os.getenv(
    "EXPORTS_DIR",
    os.path.abspath(os.path.join(os.path.dirname(__file__), "../../exports")),
)

# ---------------------------------------------------------------------------
# Organizer-native voice (Milestone R, Plan A)
#
# Powerbuilder writes in the voice of the popular-education organizing
# tradition (re:power, Wellstone, Ruckus, FWD.us Community Accelerator).
# That means three things, enforced by SYSTEM_PROMPT and _build_prompt:
#
#   1. The opening paragraph names a Theory of Change in the form
#      "If we do X, then Y will happen" before any tactics.
#   2. The body uses canonical organizing vocabulary (base, target,
#      persuadable, ladder of engagement, etc.) instead of generic
#      marketing terms (followers, reach, conversion).
#   3. A short "What this won't do" footer names the limits of the plan
#      so the user knows where human judgment still has to step in.
#
# Power type is inferred from the user query using keyword scoring (see
# _infer_power_type below). The result is surfaced both to the LLM and
# to structured_data so the chat UI can render a Power Type chip.
# ---------------------------------------------------------------------------

POWER_TYPES = {
    "through": {
        "label": "Power Through",
        "definition": "taking power via electoral or state mechanisms (winning office, winning legislation, winning the budget)",
    },
    "over": {
        "label": "Power Over",
        "definition": "forcing institutions or decision-makers to concede through public pressure, escalation, or disruption",
    },
    "with": {
        "label": "Power With",
        "definition": "building alternative structures and mutual-aid capacity inside the community itself",
    },
}

# Keyword → power type. Scored by simple substring count on the lowercased
# query plus the active_agents list. The keyword sets are intentionally short
# and high-signal: a term that points clearly at one of the three power types.
_POWER_TYPE_KEYWORDS = {
    "through": [
        "election", "electoral", "voter", "vote", "ballot", "precinct",
        "turnout", "primary", "general election", "candidate", "campaign",
        "district", "win number", "gotv", "register", "poll",
    ],
    "over": [
        "pressure", "target", "decision-maker", "decision maker",
        "protest", "march", "rally", "boycott", "escalat", "disrupt",
        "corporate", "council", "city hall", "sit-in", "strike",
        "legislator", "lobby", "demand", "accountab",
    ],
    "with": [
        "mutual aid", "co-op", "cooperative", "community fridge",
        "deportation defense", "know your rights", "care work",
        "caregiv", "tenant", "union drive", "alternative", "build base",
        "base building", "leader development", "popular education",
    ],
}

# Canonical organizing vocabulary (subset of the 30-term glossary in
# notes/organizing_pdfs_analysis.md). Injected verbatim into the user
# prompt so the LLM has the exact phrasing on hand when it writes.
ORGANIZING_GLOSSARY = (
    "- base: people already aligned with the cause, who must be activated.\n"
    "- base building: ongoing work of expanding the number of committed supporters.\n"
    "- constituency: the population whose interests we represent and organize.\n"
    "- target: the specific decision-maker with power to grant our demands.\n"
    "- theory of change: causal logic in the form 'If we do X, then Y will happen'.\n"
    "- ladder of engagement: progression from low-commitment to high-commitment actions.\n"
    "- rung: one step on the ladder of engagement.\n"
    "- persuadable: person not yet committed who could be moved with the right framing.\n"
    "- power map: visual tool plotting stakeholders by power level and alignment.\n"
    "- spectrum of allies: continuum from active support to active opposition.\n"
    "- escalation: deliberate increase in pressure or disruption across a campaign arc.\n"
    "- polarize: deliberately force people to take sides, reducing the number of bystanders.\n"
    "- PSAA: content structure of Problem, Solution, Action, Ask.\n"
    "- mobilize vs. organize: mobilize activates known supporters, organize develops new leaders.\n"
    "- declare victory and run: publish a momentum-sustaining win frame even when full demands are not met."
)

GENERIC_TO_ORGANIZER = (
    "When tempted to write 'followers' say 'base'. When tempted to write 'audience' "
    "say 'constituency' or 'persuadable audience'. When tempted to write 'engagement "
    "funnel' say 'ladder of engagement'. When tempted to write 'call to action' say "
    "'ask' and locate it on a specific rung. Never substitute marketing language for "
    "organizing language."
)


def _infer_power_type(query: str, active_agents: list | None = None) -> str:
    """
    Infer which of the three re:power power types this query is building.

    Returns one of 'through', 'over', 'with'. Falls back to 'through' on a
    tie or when nothing matches, since the most common Powerbuilder query
    is a partisan electoral run (researcher + win_number + precincts +
    messaging + cost_calculator). Scoring is intentionally tiny so it is
    cheap to run on every synthesis call and easy to reason about in tests.
    """
    haystack = (query or "").lower()
    if active_agents:
        haystack = haystack + " " + " ".join(str(a).lower() for a in active_agents)

    scores = {key: 0 for key in POWER_TYPES}
    for key, terms in _POWER_TYPE_KEYWORDS.items():
        for term in terms:
            if term in haystack:
                scores[key] += 1

    # Electoral signal from the agent list: if any of researcher,
    # election_results, win_number, precincts ran, that is a Power Through
    # signal worth one extra point so the tiebreaker leans correctly.
    if active_agents:
        agent_set = {str(a).lower() for a in active_agents}
        if agent_set & {"election_results", "win_number", "precincts"}:
            scores["through"] += 1

    best_score = max(scores.values())
    if best_score == 0:
        return "through"
    # On tie, prefer the order through > over > with (most common first).
    for key in ("through", "over", "with"):
        if scores[key] == best_score:
            return key
    return "through"


SYSTEM_PROMPT = (
    "You are a senior organizer in the popular-education tradition (re:power, "
    "Wellstone, Ruckus Society, FWD.us Community Accelerator). You have received "
    "findings from specialist analysts. Synthesize them into a single clear "
    "non-repetitive professional briefing written in first person plural "
    "('we', 'our campaign'). "
    "Methodology comes before tactics. Open every briefing with an explicit "
    "Theory of Change in the form 'If we do X, then Y will happen' before "
    "naming any tactic, channel, or content. "
    "Use canonical organizing vocabulary (base, persuadable, target, ladder "
    "of engagement, rung, spectrum of allies, escalation, mobilize vs. "
    "organize, PSAA), not generic marketing language (followers, audience "
    "reach, engagement funnel, conversion). "
    "Do not invent information not present in the inputs. "
    "Defer to the most specific and most recently dated source when findings "
    "conflict. Do not mention AI, agents, or automated tools."
)

# Agents whose presence together signals a full political plan run.
PLAN_AGENTS = {"researcher", "election_results", "win_number", "precincts", "messaging", "cost_calculator"}

# Section titles used in the plan docx — must match the template and synthesis prompt.
PLAN_SECTION_ORDER = [
    "Executive Summary",
    "District Background",
    "Target Universe and Demographics",
    "Geographic Targeting",
    "Messaging Strategy",
    "Budget Estimate",
    "Win Number Calculation",
    "Program Recommendations",
]

# ---------------------------------------------------------------------------
# Internal utilities
# ---------------------------------------------------------------------------


def _hash(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()


def _dedup(items: list) -> list:
    """Remove exact-duplicate strings by content hash, preserving insertion order."""
    seen, unique = set(), []
    for item in items:
        h = _hash(str(item))
        if h not in seen:
            seen.add(h)
            unique.append(item)
    return unique


def _most_recent_date(research_results: list) -> str:
    """
    Scan memo headers written by researcher.py for DATE fields and return
    the most recent date string found. Returns "unknown" if none parse.
    """
    dated = []
    for memo in research_results:
        for match in re.finditer(r"\| DATE: ([^|\-\n]+?) ---", memo):
            raw = match.group(1).strip()
            for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%m/%d/%Y",
                        "%B %Y", "%b %Y", "%B %d, %Y", "%Y"):
                try:
                    dated.append((datetime.strptime(raw, fmt), raw))
                    break
                except ValueError:
                    continue
    if dated:
        dated.sort(key=lambda x: x[0], reverse=True)
        return dated[0][1]
    return "unknown"


def _get_entry(structured_data: list, agent: str) -> Optional[dict]:
    """Return the first structured_data entry from the named agent, or None."""
    return next((d for d in structured_data if d.get("agent") == agent), None)


def _district_label(structured_data: list) -> str:
    """Build a human-readable district label from any agent's geographic context."""
    for agent in ("precincts", "win_number", "finance"):
        entry = _get_entry(structured_data, agent)
        if entry and entry.get("district_type") and entry.get("district_id"):
            dt  = entry["district_type"].replace("_", " ").title()
            did = entry["district_id"]
            return f"Statewide {dt}" if did == "statewide" else f"{dt} {did}"
    return "Target District"


def _safe_filename(label: str, ext: str) -> str:
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^\w]+", "_", label).strip("_").lower()
    return f"{ts}_{safe}.{ext}"


def _fmt_int(val) -> str:
    """Format an integer with commas, or return 'N/A'."""
    try:
        return f"{int(val):,}"
    except (TypeError, ValueError):
        return "N/A"


def _fmt_pct(val) -> str:
    try:
        return f"{float(val) * 100:.1f}%"
    except (TypeError, ValueError):
        return "N/A"


def _fmt_money(val) -> str:
    try:
        return f"${float(val):,.0f}"
    except (TypeError, ValueError):
        return "N/A"


# ---------------------------------------------------------------------------
# Table builders (structured_data → (headers, rows) tuples)
# ---------------------------------------------------------------------------


def _win_table(win_entry: dict) -> tuple:
    headers = ["Metric", "Value"]
    rows = [
        ["Win Number (votes needed)",  _fmt_int(win_entry.get("win_number"))],
        ["Projected Turnout",          _fmt_int(win_entry.get("projected_turnout"))],
        ["Persuadable Universe",       _fmt_int(win_entry.get("persuadable_universe"))],
        ["Voter Universe (CVAP)",      _fmt_int(win_entry.get("voter_universe_cvap"))],
        ["Avg Historical Turnout",     _fmt_pct(win_entry.get("avg_turnout_pct"))],
        ["Victory Margin Target",      _fmt_pct(win_entry.get("victory_margin"))],
        ["Historical Basis",           str(win_entry.get("historical_context", "N/A"))],
    ]
    return headers, rows


def _precinct_table(precincts: list) -> tuple:
    """Build (headers, rows) from a list of precinct dicts."""
    if not precincts:
        return [], []
    exclude = {"precinct_geoid", "approximate_boundary", "precinct_name"}
    metric_keys = [k for k in precincts[0] if k not in exclude]
    headers = ["Precinct"] + [k.replace("_", " ").title() for k in metric_keys]
    rows = []
    for p in precincts:
        row = [p.get("precinct_name") or p.get("precinct_geoid", "")]
        for k in metric_keys:
            val = p.get(k)
            row.append(f"{float(val):,.0f}" if isinstance(val, (int, float)) else str(val or ""))
        rows.append(row)
    return headers, rows


def _budget_tables(finance_entry: dict) -> tuple:
    """
    Returns two (headers, rows) tuples:
      - unit_cost_table: per-contact rates, with contact goals if budget was provided
      - fec_table:       historical FEC category breakdown, or None if unavailable
    """
    unit_costs     = finance_entry.get("unit_costs", {})
    budget_program = finance_entry.get("budget_program")
    full_est       = finance_entry.get("full_program_estimate")

    tactic_labels = {
        "door_knock":   "Door Knock",
        "phone_call":   "Phone Call",
        "text_message": "Text Message",
        "mail_piece":   "Mail Piece",
        "digital":      "Digital Advertising",
    }

    if budget_program:
        h = ["Tactic", "Unit Cost", "Contact Goal", "Budget Allocated"]
        r = [
            [
                tactic_labels.get(t, t),
                f"${d['unit_cost']:.2f}",
                f"{d['contacts']:,}",
                _fmt_money(d["budget_allocated"]),
            ]
            for t, d in budget_program.items()
        ]
    else:
        h = ["Tactic", "Unit Cost"]
        r = [[tactic_labels.get(k, k), f"${v:.2f}"] for k, v in unit_costs.items()]

    unit_cost_table = (h, r)

    fec_table = None
    if full_est and full_est.get("total"):
        fh = ["Category", "Estimated Spend"]
        fr = [
            ["Total Program (FEC average)",  _fmt_money(full_est.get("total"))],
            ["Personnel",                    _fmt_money(full_est.get("personnel", 0))],
            ["Mail",                         _fmt_money(full_est.get("mail", 0))],
            ["Digital",                      _fmt_money(full_est.get("digital", 0))],
            ["Phones",                       _fmt_money(full_est.get("phones", 0))],
            ["Miscellaneous",               _fmt_money(full_est.get("miscellaneous", 0))],
        ]
        fec_table = (fh, fr)

    return unit_cost_table, fec_table


# ---------------------------------------------------------------------------
# Synthesis (LLM)
# ---------------------------------------------------------------------------


_ATTRIBUTION = (
    "*Research sourced from American Bridge Research Books, Analyst Institute, "
    "CIRCLE, and Powerbuilder's curated research corpus.*"
)


def _build_prompt(
    query: str,
    research_context: str,
    structured_context: str,
    active_agents: list,
    errors: list,
    is_plan: bool,
    district_label: str,
    power_type: str = "through",
) -> str:

    error_block = ""
    if errors:
        error_block = (
            "\n⚠️ The following agents encountered issues and their outputs may be "
            "incomplete, note this where relevant:\n"
            + "\n".join(f"  - {e}" for e in errors)
            + "\n"
        )

    pt = POWER_TYPES.get(power_type, POWER_TYPES["through"])
    power_block = (
        f"\nINFERRED POWER TYPE: {pt['label']} ({pt['definition']}).\n"
        "Frame the Theory of Change in the opening paragraph against this power type "
        "and reference it once more by name when explaining strategy.\n"
    )

    glossary_block = (
        "\nCANONICAL ORGANIZING VOCABULARY (use these terms over generic marketing language):\n"
        f"{ORGANIZING_GLOSSARY}\n\n{GENERIC_TO_ORGANIZER}\n"
    )

    wont_do_block = (
        "\nAfter the main content, append a final H2 section titled exactly "
        "'What This Won\u2019t Do' (use a curly apostrophe). In 2 to 4 short bullets, "
        "name the limits of this plan that still require human judgment, for example: "
        "this plan does not replace 1:1 conversations with the base, this plan does not "
        "include voter file or VAN access, this plan assumes the named target is the "
        "correct decision-maker (verify before escalation), this plan does not handle "
        "rapid-response decisions in real time. Pick limits that actually fit the briefing.\n"
    )

    if is_plan:
        structure = f"""
Produce a complete Political Program Plan in Markdown for {district_label}.
Use H1 for the document title, H2 for each section. Do NOT use markdown tables —
structured data tables will be inserted programmatically. Use bullet lists and bold.

Required H2 sections (use these exact titles):

## Executive Summary
2–3 paragraphs on the race, the opportunity, and the strategy to win.

## District Background
Describe the district: geography, jurisdiction type, historical partisan lean,
key communities, and any relevant political context. Draw from research findings.

## Target Universe and Demographics
Summarise the demographic patterns across the target precincts in narrative form —
age, race/ethnicity, CVAP composition, and voter registration trends.
Do not reproduce raw numbers in a table — the precinct table will be inserted after this section.

## Geographic Targeting
Describe the logic behind the target precinct selection and how geographic
concentration serves the win-number goal. Reference specific precincts where notable.

## Messaging Strategy
Present the messaging strategy: summarise the canvassing approach, phone banking plan,
text messaging goals, mail narrative themes, and digital ad strategy, then include
the full text of all scripts and copy produced by the messaging analyst.

## Budget Estimate
Narrative interpretation of the budget analysis. If a specific budget was provided,
discuss what program it funds and what trade-offs were made. Do not reproduce tables —
they will be inserted after this section.

## Win Number Calculation
Narrative interpretation of the win number: projected turnout, CVAP universe,
historical turnout context, and what the number means for program scale.
Do not reproduce raw numbers in a table — the win number table will be inserted here.

## Program Recommendations
3–5 concrete, prioritised action items for the campaign with rationale for each.
Name each as a rung on the ladder of engagement where it fits.

{wont_do_block}
Then end with (italicised):
{_ATTRIBUTION}
"""
    else:
        structure = f"""
Produce a professional organizer briefing in Markdown responding to: "{query}"
Use H2 for major sections. Use bullet lists and bold for emphasis.
The first paragraph must state the Theory of Change in the form
"If we do X, then Y will happen" before any tactics.
{wont_do_block}
Then end with: {_ATTRIBUTION}
"""

    return f"""{error_block}{power_block}{glossary_block}
USER REQUEST: {query}
AGENTS THAT CONTRIBUTED: {', '.join(active_agents) if active_agents else 'none'}

RESEARCH FINDINGS:
{research_context}

STRUCTURED DATA (win number, precincts, budget):
{structured_context}

{structure}"""


def _synthesize(state: AgentState, is_plan: bool) -> str:
    """Single GPT-4o call with the senior strategist system prompt."""
    research_results = _dedup(state.get("research_results", []))
    structured_data  = state.get("structured_data", [])
    active_agents    = state.get("active_agents", [])
    errors           = state.get("errors", [])

    research_ctx   = "\n\n".join(research_results) or "No research collected."
    structured_ctx = str(structured_data) if structured_data else "No structured data collected."
    district_lbl   = _district_label(structured_data)

    power_type = _infer_power_type(state.get("query", ""), active_agents)

    prompt = _build_prompt(
        query=state.get("query", ""),
        research_context=research_ctx,
        structured_context=structured_ctx,
        active_agents=active_agents,
        errors=errors,
        is_plan=is_plan,
        district_label=district_lbl,
        power_type=power_type,
    )

    llm = get_completion_client(temperature=0.3)
    return llm.invoke(
        [{"role": "system", "content": SYSTEM_PROMPT},
         {"role": "user",   "content": prompt}]
    ).content


# ---------------------------------------------------------------------------
# Markdown section parser (for docx)
# ---------------------------------------------------------------------------


def _parse_sections(text: str) -> dict:
    """
    Split Markdown on `## ` headers into {title: body} dict.
    Content before the first ## header is stored under "preamble".
    """
    sections: dict = {}
    current_title = "preamble"
    current_lines: list = []
    for line in text.splitlines():
        if line.startswith("## "):
            sections[current_title] = "\n".join(current_lines).strip()
            current_title = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    sections[current_title] = "\n".join(current_lines).strip()
    return sections


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------


def _bold_runs(para, text: str):
    """Add runs to a paragraph, bolding **text** spans."""
    for i, part in enumerate(re.split(r"\*\*([^*]+)\*\*", text)):
        if part:
            run = para.add_run(part)
            run.bold = bool(i % 2)


def _add_prose(doc, text: str):
    """
    Add Markdown prose to a docx Document.
    Handles: ### headings, - bullets, *italic recency notes*, bold (**text**).
    Skips markdown table rows (lines starting with |) — tables added separately.
    """
    for line in text.splitlines():
        s = line.strip()
        if not s or s.startswith("|") or re.match(r"^[-|: ]+$", s):
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            _bold_runs(doc.add_paragraph(style="List Bullet"), s[2:])
        elif s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            run = doc.add_paragraph().add_run(s.strip("*"))
            run.italic = True
        else:
            _bold_runs(doc.add_paragraph(), s)


def _add_table(doc, headers: list, rows: list):
    """Add a styled data table to a docx Document."""
    if not rows:
        return
    from docx.shared import Pt, RGBColor

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = str(val)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Paid-media section: styled DOCX renderer for the structured paid_media dict
# ---------------------------------------------------------------------------


def _strip_inline_paid_media(text: str) -> str:
    """
    Remove the `### Paid Media Plan` block from a Markdown chunk.

    The narrative formatter in chat/agents/paid_media.py emits a full markdown
    block, including a pipe-delimited table, that the LLM may keep verbatim
    inside the Budget Estimate section. _add_prose drops table rows silently,
    so leaving the block in would yield a half-rendered duplicate. We strip
    everything from the `### Paid Media Plan` heading through the next
    `## ` (next H2) or `### ` (next H3) or end of text — whichever comes first.
    """
    if "### Paid Media Plan" not in text:
        return text
    lines = text.splitlines()
    out: list = []
    skipping = False
    for line in lines:
        s = line.strip()
        if not skipping and s.startswith("### Paid Media Plan"):
            skipping = True
            continue
        if skipping:
            # Stop skipping at the next H2 or sibling H3.
            if s.startswith("## ") or (s.startswith("### ") and not s.startswith("### Paid Media Plan")):
                skipping = False
                out.append(line)
        else:
            out.append(line)
    return "\n".join(out)


def _paid_media_digital_table(channels: list) -> tuple:
    """Build (headers, rows) for the paid-media digital channel table."""
    headers = [
        "Channel", "Spend", "Share", "CPM (low: high)",
        "Impressions (mid)", "Est. Reach", "Est. Lift (pp)",
    ]
    rows = []
    for c in channels:
        cpm = f"${c['cpm_low']:.0f}: ${c['cpm_high']:.0f}"
        imp_mid = f"{c['impressions']['mid']:,}"
        reach   = f"{c['reach']:,}" if c.get("reach") else "n/a"
        if c.get("saturated"):
            reach += " (capped)"
        if c.get("points_lift"):
            lift = f"{c['points_lift']['low']:.2f}: {c['points_lift']['high']:.2f}"
        else:
            lift = "n/a"
        rows.append([
            c["label"],
            _fmt_money(c["spend"]),
            f"{c['share_pct']:.0f}%",
            cpm,
            imp_mid,
            reach,
            lift,
        ])
    return headers, rows


def _render_paid_media_section(doc, paid_media: dict):
    """
    Render the paid_media dict from finance structured_data as a styled DOCX
    subsection under the Budget Estimate H2. Uses _add_table for the digital
    channel rollup, bullets for non-digital allocation, a paragraph for the
    total persuasion lift, bullets for notes, and an italic source citation.
    """
    if not paid_media:
        return

    doc.add_heading("Paid Media Plan", level=3)

    # Lead paragraph: tier, totals, in-language note.
    district_bit = (
        f" for {paid_media['district_label']}" if paid_media.get("district_label") else ""
    )
    lead = (
        f"Tier: **{paid_media['tier_label']}**. Total program budget "
        f"**{_fmt_money(paid_media['budget'])}**{district_bit}, of which "
        f"**{_fmt_money(paid_media['digital_spend_total'])}** is paid digital and "
        f"**{_fmt_money(paid_media['non_digital_spend_total'])}** is SMS and print."
    )
    _bold_runs(doc.add_paragraph(), lead)

    if paid_media.get("in_language_pricing"):
        _bold_runs(
            doc.add_paragraph(),
            f"In-language pricing applied for **{paid_media['language_intent']}**: "
            f"22.5 percent under English CPM in the same DMA.",
        )

    # Digital channels table.
    if paid_media.get("channels"):
        _bold_runs(doc.add_paragraph(), "**Digital channels**")
        h, r = _paid_media_digital_table(paid_media["channels"])
        _add_table(doc, h, r)

    # Non-digital allocation as bullets.
    if paid_media.get("non_digital"):
        _bold_runs(doc.add_paragraph(), "**Non-digital allocation (in same total)**")
        for n in paid_media["non_digital"]:
            _bold_runs(
                doc.add_paragraph(style="List Bullet"),
                f"{n['label']}: {_fmt_money(n['spend'])} ({n['share_pct']:.0f}%)",
            )

    # Total persuasion lift summary.
    lift_low  = paid_media.get("total_points_lift_low")  or 0
    lift_high = paid_media.get("total_points_lift_high") or 0
    if lift_low or lift_high:
        _bold_runs(
            doc.add_paragraph(),
            f"**Estimated total persuasion lift in target universe:** "
            f"{lift_low:.2f} to {lift_high:.2f} percentage points "
            f"(digital channels only; door knocks and direct mail are tracked separately "
            f"in the operational program above).",
        )

    # Notes bullets.
    if paid_media.get("notes"):
        _bold_runs(doc.add_paragraph(), "**Notes**")
        for note in paid_media["notes"]:
            _bold_runs(doc.add_paragraph(style="List Bullet"), note)

    # Italic source citation.
    source = paid_media.get("source") or (
        "Powerbuilder corpus file 07: paid-media digital benchmarks"
    )
    src_para = doc.add_paragraph()
    src_run  = src_para.add_run(f"Source: {source}.")
    src_run.italic = True

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Format handlers — unified signature: (synthesis, state, district_label)
# _write_text ignores state and district_label via **_
# ---------------------------------------------------------------------------


def _write_text(synthesis: str, **_) -> dict:
    """Text / Markdown: write synthesis directly to final_answer."""
    return {"final_answer": synthesis}


def _write_docx(synthesis: str, state: AgentState, district_label: str) -> dict:
    """
    Build a Word document. Prose from LLM synthesis, tables from structured_data.
    """
    try:
        from docx import Document
    except ImportError:
        return {
            "final_answer": synthesis,
            "errors": ["ExportAgent: python-docx is not installed (pip install python-docx)."],
        }

    structured_data = state.get("structured_data", [])
    win_entry       = _get_entry(structured_data, "win_number")
    precinct_entry  = _get_entry(structured_data, "precincts")
    finance_entry   = _get_entry(structured_data, "finance")
    precincts       = (precinct_entry or {}).get("precincts", [])

    # Open the branded template if it exists; fall back to a blank document.
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
        # Clear body content (placeholder text) while preserving styles and footer.
        from docx.oxml.ns import qn
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    else:
        doc = Document()
    doc.add_heading(f"{district_label} — Political Program Plan", 0)

    # If the finance entry contains a structured paid_media plan, the styled
    # renderer below will draw it from the dict. Strip the inline markdown
    # version from the synthesis so we don't get a half-rendered duplicate
    # (the markdown table rows would be silently dropped by _add_prose).
    paid_media_plan = (finance_entry or {}).get("paid_media") if finance_entry else None
    synthesis_for_docx = _strip_inline_paid_media(synthesis) if paid_media_plan else synthesis

    sections = _parse_sections(synthesis_for_docx)

    # Use PLAN_SECTION_ORDER when this is a plan doc; otherwise use whatever the LLM produced.
    ordered_titles = (
        PLAN_SECTION_ORDER
        if any(t in sections for t in PLAN_SECTION_ORDER)
        else [k for k in sections if k != "preamble"]
    )

    for title in ordered_titles:
        section_text = sections.get(title, "")
        doc.add_heading(title, level=2)
        if section_text:
            _add_prose(doc, section_text)

        # Inject data tables after the matching section
        if "Win Number" in title:
            if win_entry:
                doc.add_heading("Win Number Summary", level=3)
                _add_table(doc, *_win_table(win_entry))

        elif "Geographic" in title:
            if precincts:
                doc.add_heading(f"Top {len(precincts)} Target Precincts", level=3)
                h, r = _precinct_table(precincts)
                _add_table(doc, h, r)

        elif "Budget" in title:
            if finance_entry:
                (h, r), fec = _budget_tables(finance_entry)
                doc.add_heading("Per-Contact Rate Estimates", level=3)
                _add_table(doc, h, r)
                if fec:
                    doc.add_heading("Historical FEC Spending (comparable cycles)", level=3)
                    _add_table(doc, *fec)
                if paid_media_plan:
                    _render_paid_media_section(doc, paid_media_plan)

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    path = os.path.join(EXPORTS_DIR, _safe_filename(district_label, "docx"))
    doc.save(path)
    logger.info(f"ExportAgent: saved docx → {path}")

    return {"final_answer": synthesis, "generated_file_path": path}


def _write_xlsx(synthesis: str, state: AgentState, district_label: str) -> dict:
    """
    Build an Excel workbook with three sheets:
      Sheet 1 — Precinct Targets (from precincts structured_data)
      Sheet 2 — Win Number Summary
      Sheet 3 — Budget Estimate
    """
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return {
            "final_answer": synthesis,
            "errors": ["ExportAgent: openpyxl is not installed (pip install openpyxl)."],
        }

    structured_data = state.get("structured_data", [])
    win_entry       = _get_entry(structured_data, "win_number")
    precinct_entry  = _get_entry(structured_data, "precincts")
    finance_entry   = _get_entry(structured_data, "finance")
    precincts       = (precinct_entry or {}).get("precincts", [])

    wb = openpyxl.Workbook()

    HEADER_FILL  = PatternFill("solid", fgColor="1F4E79")
    HEADER_FONT  = Font(color="FFFFFF", bold=True)
    HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def _style_headers(ws, n_cols: int):
        for cell in ws[1]:
            cell.fill      = HEADER_FILL
            cell.font      = HEADER_FONT
            cell.alignment = HEADER_ALIGN
        for i in range(1, n_cols + 1):
            ws.column_dimensions[get_column_letter(i)].width = 20

    # ---- Sheet 1: Precinct Targets ----
    ws1 = wb.active
    ws1.title = "Precinct Targets"
    if precincts:
        h, rows = _precinct_table(precincts)
        ws1.append(h)
        for row in rows:
            ws1.append(row)
        _style_headers(ws1, len(h))
    else:
        ws1.append(["No precinct data available."])

    # ---- Sheet 2: Win Number ----
    ws2 = wb.create_sheet("Win Number")
    h2 = ["Metric", "Value"]
    ws2.append(h2)
    if win_entry:
        _, rows = _win_table(win_entry)
        for row in rows:
            ws2.append(row)
    else:
        ws2.append(["No win number data available.", ""])
    _style_headers(ws2, 2)
    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 20

    # ---- Sheet 3: Budget Estimate ----
    ws3 = wb.create_sheet("Budget Estimate")
    if finance_entry:
        (h3, rows3), fec = _budget_tables(finance_entry)
        ws3.append(h3)
        for row in rows3:
            ws3.append(row)
        _style_headers(ws3, len(h3))

        if fec:
            fec_h, fec_rows = fec
            ws3.append([])  # blank spacer row
            ws3.append(fec_h)
            fec_header_row = ws3.max_row
            for row in fec_rows:
                ws3.append(row)
            # Style the FEC sub-header
            for cell in ws3[fec_header_row]:
                cell.fill      = HEADER_FILL
                cell.font      = HEADER_FONT
                cell.alignment = HEADER_ALIGN
    else:
        ws3.append(["No budget data available."])

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    path = os.path.join(EXPORTS_DIR, _safe_filename(district_label, "xlsx"))
    wb.save(path)
    logger.info(f"ExportAgent: saved xlsx → {path}")

    brief = (
        f"Excel workbook generated with precinct targets, win number, and budget data "
        f"for {district_label}.\nFile saved to: {path}"
    )
    return {"final_answer": brief, "generated_file_path": path}


def _write_csv(synthesis: str, state: AgentState, district_label: str) -> dict:
    """
    Flat CSV of target precincts. Falls back to a structured_data summary
    (agent, key, value rows) if no precinct data is available.
    """
    structured_data = state.get("structured_data", [])
    precinct_entry  = _get_entry(structured_data, "precincts")
    precincts       = (precinct_entry or {}).get("precincts", [])

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    path = os.path.join(EXPORTS_DIR, _safe_filename(district_label, "csv"))

    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv_module.writer(f)
        if precincts:
            h, rows = _precinct_table(precincts)
            writer.writerow(h)
            writer.writerows(rows)
        else:
            # Fallback: flat dump of all non-nested structured_data fields
            writer.writerow(["agent", "key", "value"])
            for entry in structured_data:
                agent_name = entry.get("agent", "unknown")
                for k, v in entry.items():
                    if k != "agent" and not isinstance(v, (list, dict)):
                        writer.writerow([agent_name, k, v])

    logger.info(f"ExportAgent: saved csv → {path}")

    brief = (
        f"{len(precincts)} target precincts exported to CSV for {district_label}.\n"
        f"File saved to: {path}"
    )
    return {"final_answer": brief, "generated_file_path": path}


# ---------------------------------------------------------------------------
# Format handler dispatch
# ---------------------------------------------------------------------------

_HANDLERS = {
    "text":     _write_text,
    "markdown": _write_text,
    "docx":     _write_docx,
    "xlsx":     _write_xlsx,
    "csv":      _write_csv,
}


# ---------------------------------------------------------------------------
# Node
# ---------------------------------------------------------------------------


def export_node(state: AgentState) -> dict:
    """
    Final LangGraph node. Synthesizes all agent outputs and formats the result.

    Reads from state:
      query, research_results, structured_data, active_agents, errors, output_format

    Writes to state:
      final_answer         — markdown synthesis, always set
      generated_file_path  — set for full plans (docx) and explicit csv/xlsx requests
      errors               — format failures (non-fatal)

    Full plan behaviour (is_plan=True):
      Always calls _write_docx regardless of output_format. _write_docx returns both
      final_answer (the markdown synthesis for chat display) and generated_file_path
      (the saved .docx path for the download button). output_format is ignored so that
      the router's MARKDOWN default does not suppress file generation.

    Non-plan behaviour (is_plan=False):
      Uses the router-specified output_format as before.
    """
    active_agents = state.get("active_agents", [])
    output_format = state.get("output_format", "text")
    structured_data = state.get("structured_data", [])

    is_plan      = PLAN_AGENTS.issubset(set(active_agents))
    district_lbl = _district_label(structured_data)

    # Compute the inferred power type up front so it is available both to
    # the synthesizer prompt and to the structured_data the UI reads. Note:
    # AgentState.structured_data uses operator.add as its reducer, so we
    # return ONLY the new power_type entry from this node and let LangGraph
    # merge it into the running list.
    inferred_power = _infer_power_type(state.get("query", ""), active_agents)
    power_meta = {
        "agent": "power_type",
        "power_type": inferred_power,
        "label": POWER_TYPES[inferred_power]["label"],
        "definition": POWER_TYPES[inferred_power]["definition"],
    }

    # -----------------------------------------------------------------------
    # 1. Synthesize
    # -----------------------------------------------------------------------
    try:
        synthesis = _synthesize(state, is_plan)
    except Exception as e:
        logger.error(f"ExportAgent: synthesis LLM call failed — {e}")
        fallback = "\n\n".join(_dedup(state.get("research_results", [])))
        return {
            "final_answer": fallback or "Synthesis failed. Check the errors field.",
            "errors":       [f"ExportAgent: LLM synthesis failed — {e}"],
        }

    # -----------------------------------------------------------------------
    # 2. Format
    # Full plans always use _write_docx so the Word file is generated alongside
    # the markdown final_answer. Non-plan queries use the router-specified format.
    # -----------------------------------------------------------------------
    handler = _write_docx if is_plan else _HANDLERS.get(output_format, _write_text)
    try:
        result = handler(synthesis, state=state, district_label=district_lbl)
    except Exception as e:
        logger.error(f"ExportAgent: formatting failed (primary) - {e}")
        result = {
            "final_answer": synthesis,
            "errors": [f"ExportAgent: Could not generate output file - {e}"],
        }

    # Track all generated files so the UI can offer multiple downloads.
    generated_files: list[str] = []
    primary = result.get("generated_file_path")
    if primary:
        generated_files.append(primary)

    # ---- 2b. CSV companion for full plans ----------------------------------
    # Plans typically include a target list of precincts. Always emit a CSV
    # alongside the DOCX so the operator gets both the script (DOCX) and the
    # walk/call list (CSV) from a single chat turn. Non-plan queries that
    # explicitly asked for csv already used _write_csv as the primary handler
    # (skip in that case).
    if is_plan:
        precinct_entry = _get_entry(structured_data, "precincts")
        if precinct_entry and precinct_entry.get("precincts"):
            try:
                csv_result = _write_csv(synthesis, state=state, district_label=district_lbl)
                csv_path = csv_result.get("generated_file_path")
                if csv_path and csv_path not in generated_files:
                    generated_files.append(csv_path)
            except Exception as e:
                logger.error(f"ExportAgent: CSV companion failed - {e}")
                existing_errors = list(result.get("errors", []))
                existing_errors.append(
                    f"ExportAgent: Could not generate CSV companion - {e}"
                )
                result["errors"] = existing_errors

    if generated_files:
        result["generated_files"] = generated_files

    # Append the power_type entry to structured_data via the operator.add
    # reducer so the chat UI and tests can read it without having to call
    # _infer_power_type again.
    result["structured_data"] = [power_meta]

    logger.info(
        f"ExportAgent: format={'docx(plan)' if is_plan else output_format} | "
        f"district={district_lbl} | power={inferred_power} | files={generated_files or 'none'}"
    )
    return result
